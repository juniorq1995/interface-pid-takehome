"""Generate synthetic P&ID + SOP fixtures for development and tests.

These are NOT the real Interface assignment data — placeholders only, used to
prove the pipeline runs end-to-end before the real diagram.pdf / sop.docx
arrive. Swap them out at data/pid/diagram.pdf and data/sop/sop.docx.

Scene (by design, to exercise every cross-reference finding type):
  T-101 (tank) -- P-101 (pump) -- V-101 (valve) -- FT-101 (instrument)
  V-102 (valve, isolated)       -- never mentioned in the SOP   -> MISSING_IN_SOP
  unlabeled tank glyph          -- no readable tag              -> UNRESOLVED_TAG
  SOP mentions PSV-201          -- not drawn in the diagram      -> MISSING_IN_PID
  SOP calls T-101 a "pump"      -- diagram tag prefix says tank  -> TYPE_MISMATCH
  SOP says "T-101 connects to FT-101" -- no such edge exists     -> CONNECTION_MISMATCH
  SOP says "P-101 connects to V-101"  -- matches a real edge     -> correctly NOT flagged
"""
from __future__ import annotations

from pathlib import Path

import cv2
import docx
import numpy as np
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
PID_PATH = ROOT / "data" / "pid" / "diagram.pdf"
SOP_PATH = ROOT / "data" / "sop" / "sop.docx"

BLACK = (0, 0, 0)
WHITE = (255, 255, 255)


def _label(img: np.ndarray, text: str, cx: int, bottom_y: int) -> None:
    (tw, th), _ = cv2.getTextSize(text, cv2.FONT_HERSHEY_SIMPLEX, 0.9, 2)
    cv2.putText(img, text, (cx - tw // 2, bottom_y + th + 20), cv2.FONT_HERSHEY_SIMPLEX, 0.9, BLACK, 2)


def _draw_tank(img: np.ndarray, cx: int, cy: int) -> int:
    x0, y0, x1, y1 = cx - 60, cy - 80, cx + 60, cy + 80
    cv2.rectangle(img, (x0, y0), (x1, y1), BLACK, -1)
    return y1


def _draw_pump(img: np.ndarray, cx: int, cy: int, radius: int = 45) -> int:
    cv2.circle(img, (cx, cy), radius, BLACK, -1)
    triangle = np.array([[cx - 18, cy + 14], [cx + 18, cy + 14], [cx, cy - 20]])
    cv2.fillPoly(img, [triangle], WHITE)  # impeller cut-out -> creates a child contour
    return cy + radius


def _draw_instrument(img: np.ndarray, cx: int, cy: int, radius: int = 35) -> int:
    cv2.circle(img, (cx, cy), radius, BLACK, -1)
    return cy + radius


def _draw_valve(img: np.ndarray, cx: int, cy: int) -> int:
    """Stylized (non-ISA-literal) valve glyph: a notched rectangle. Deliberately
    simple/concave so the solidity-based classifier in shape_detection.py reliably
    tags it as 'bowtie' without depending on fragile touching-triangle rasterization."""
    w, h, notch = 45, 30, 18
    pts = np.array(
        [
            [cx - w, cy - h], [cx - notch, cy - h], [cx, cy], [cx + notch, cy - h],
            [cx + w, cy - h], [cx + w, cy + h], [cx - w, cy + h],
        ]
    )
    cv2.fillPoly(img, [pts], BLACK)
    return cy + h


def generate_pid() -> None:
    img = np.full((650, 1400, 3), 255, dtype=np.uint8)

    y = 300
    bottom = _draw_tank(img, 150, y)
    _label(img, "T-101", 150, bottom)

    bottom = _draw_pump(img, 400, y)
    _label(img, "P-101", 400, bottom)

    bottom = _draw_valve(img, 650, y)
    _label(img, "V-101", 650, bottom)

    bottom = _draw_instrument(img, 900, y)
    _label(img, "FT-101", 900, bottom)

    bottom = _draw_valve(img, 1150, y)
    _label(img, "V-102", 1150, bottom)

    # Unlabeled tank glyph -> should end up UNRESOLVED_TAG (no OCR-readable label nearby).
    cv2.rectangle(img, (670, 100), (730, 150), BLACK, -1)

    for (x1, x2) in [(150, 400), (400, 650), (650, 900)]:
        cv2.line(img, (x1, y), (x2, y), BLACK, 3)

    PID_PATH.parent.mkdir(parents=True, exist_ok=True)
    Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB)).save(PID_PATH, "PDF", resolution=200.0)
    print(f"Wrote synthetic P&ID -> {PID_PATH}")


def generate_sop() -> None:
    document = docx.Document()
    document.add_heading("Standard Operating Procedure: Feed Line Startup", level=1)
    for line in [
        "1. Confirm pump T-101 reaches operating pressure before opening V-101.",
        "2. Open valve V-101 slowly to begin flow to pump P-101.",
        "3. P-101 connects to V-101 downstream of the tank discharge.",
        "4. Flow transmitter FT-101 monitors line flow rate after V-101.",
        "5. T-101 connects to FT-101 for high-level interlock cross-check.",
        "6. PSV-201 safety valve must be inspected monthly per code requirements.",
    ]:
        document.add_paragraph(line)

    SOP_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(SOP_PATH)
    print(f"Wrote synthetic SOP -> {SOP_PATH}")


if __name__ == "__main__":
    generate_pid()
    generate_sop()

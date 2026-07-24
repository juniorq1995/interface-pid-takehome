"""Parse an SOP .docx into plain text and numbered/step paragraphs."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import docx


@dataclass
class SopDocument:
    full_text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]  # list of tables, each a list of rows, each a list of cell strings


def parse_sop(docx_path: str | Path) -> SopDocument:
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"SOP document not found: {docx_path}")

    document = docx.Document(str(docx_path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = [
        [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    full_text = "\n".join(paragraphs)
    return SopDocument(full_text=full_text, paragraphs=paragraphs, tables=tables)
